import streamlit as st
import os
import json
import requests
import numpy as np
import faiss
import time
import pickle
from github import Github, GithubException, Auth
from pathlib import Path
from io import BytesIO
from typing import List, Dict, Any, Tuple, Optional
from PyPDF2 import PdfReader
from http import HTTPStatus
import dashscope
from dashscope import TextEmbedding
from openai import OpenAI
from docx import Document
import matplotlib.pyplot as plt
from scipy.interpolate import make_interp_spline
from graphrag_retriever import GraphRAGRetriever
import base64

from basic_case_process import basic_case_process
from supplementary_case_process import supplementary_case_process
from finetune_data_process import finetune_data_process


# ==========================================
# [SECTION 0] 基础配置、样式与路径定义
# ==========================================

st.set_page_config(
    page_title="茶饮六因子AI评分器 Pro",
    page_icon="🍵",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 全局 CSS 样式
st.markdown("""
    <style>
    .main-title {font-size: 2.5em; font-weight: bold; text-align: center; color: #2E7D32; margin-bottom: 0.5em;}
    .slogan {font-size: 1.2em; font-style: italic; text-align: center; color: #558B2F; margin-bottom: 30px; font-family: "KaiTi", "楷体", serif;}
    .factor-card {background-color: #F1F8E9; padding: 15px; border-radius: 10px; margin-bottom: 10px; border-left: 5px solid #4CAF50;}
    .score-header {display:flex; justify-content:space-between; font-weight:bold; color:#2E7D32;}
    .advice-tag {font-size: 0.85em; padding: 2px 6px; border-radius: 4px; margin-top: 5px; background-color: #fff; border: 1px dashed #4CAF50; color: #388E3C; display: inline-block;}
    .master-comment {background-color: #FFFDE7; border: 1px solid #FFF9C4; padding: 15px; border-radius: 8px; font-family: "KaiTi", serif; font-size: 1.1em; color: #5D4037; margin-bottom: 20px; line-height: 1.6;}
    .ft-card {border: 1px solid #ddd; padding: 15px; border-radius: 8px; background-color: #f8f9fa; margin-top: 10px;}
    .case-card {border: 1px solid #e0e0e0; padding: 12px; border-radius: 8px; margin-bottom: 10px; background-color: #fafafa;}
    </style>
""", unsafe_allow_html=True)


class PathConfig:
    """路径管理类 —— 集中定义所有文件路径"""
    # 外部资源文件
    SRC_SYS_PROMPT = Path("sys_p.txt")
    # 运行时数据目录
    DATA_DIR = Path("./tea_data")
    RAG_DIR = Path("./tea_data/RAG")
    BACKUP_DIR = Path("./tea_backup")

    def __init__(self):
        self.DATA_DIR.mkdir(exist_ok=True)
        self.RAG_DIR.mkdir(exist_ok=True)
        self.BACKUP_DIR.mkdir(exist_ok=True)
        self.GRAPHRAG_DIR = self.DATA_DIR / "graphrag_artifacts"
        self.GRAPHRAG_DIR.mkdir(exist_ok=True)

        # --- 知识库（RAG） ---
        self.kb_index = self.DATA_DIR / "kb.index"
        self.kb_chunks = self.DATA_DIR / "kb_chunks.pkl"
        self.kb_files = self.DATA_DIR / "kb_files.json"

        # --- 判例库（基础 + 进阶） ---
        self.basic_case_data = self.DATA_DIR / "basic_case.json"
        self.supp_case_index = self.DATA_DIR / "supp_cases.index"
        self.supp_case_data = self.DATA_DIR / "supplementary_case.json"

        # --- 微调与 Prompt ---
        self.training_file = self.DATA_DIR / "deepseek_finetune.jsonl"
        self.ft_status = self.DATA_DIR / "ft_status.json"
        self.prompt_config_file = self.DATA_DIR / "prompts.json"

        # --- 模板与默认配置（位于 tea_backup） ---
        self.template_file = self.BACKUP_DIR / "template.xlsx"
        self.default_prompts = self.BACKUP_DIR / "default_prompts.json"


PATHS = PathConfig()

# 用户 Prompt 模板（含基础判例与进阶判例占位符）
DEFAULT_USER_TEMPLATE = """【待评分产品】
{product_desc}

【参考标准（知识库）】
{context_text}

【基础判例（全部）】
{basic_case_text}

【相似判例得分参考（进阶判例库）】
{case_text}

请严格输出以下JSON格式（不含Markdown）：
{{
  "master_comment": "约100字的宗师级总评，富含东方美学与文化意蕴...",
  "scores": {{
    "优雅性": {{"score": 1-9, "comment": "...", "suggestion": "..."}},
    "辨识度": {{"score": 1-9, "comment": "...", "suggestion": "..."}},
    "协调性": {{"score": 1-9, "comment": "...", "suggestion": "..."}},
    "饱和度": {{"score": 1-9, "comment": "...", "suggestion": "..."}},
    "持久性": {{"score": 1-9, "comment": "...", "suggestion": "..."}},
    "苦涩度": {{"score": 1-9, "comment": "...", "suggestion": "..."}}
  }}
}}"""


# ==========================================
# [SECTION 1] 资源与数据管理 (ResourceManager)
# ==========================================

class ResourceManager:
    """负责外部文件加载、数据持久化、微调数据管理"""

    # ---------- 通用文件读写 ----------
    @staticmethod
    def load_external_text(path: Path, fallback: str = "") -> str:
        """读取外部文本文件"""
        if path.exists():
            try:
                with open(path, "r", encoding="utf-8") as f:
                    return f.read().strip()
            except Exception as e:
                st.error(f"加载文件 {path} 失败: {e}")
        return fallback

    @staticmethod
    def load_external_json(path: Path, fallback: Any = None) -> Any:
        """读取外部 JSON 文件"""
        if path.exists():
            try:
                with open(path, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                st.error(f"加载文件 {path} 失败: {e}")
        return fallback if fallback is not None else []

    @staticmethod
    def save_json(data: Any, path: Path):
        """保存数据为 JSON 文件"""
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    # ---------- FAISS 索引 + 数据 ----------
    @staticmethod
    def save(index: Any, data: Any, idx_path: Path, data_path: Path, is_json: bool = False):
        """保存 FAISS 索引和对应的数据文件"""
        if index:
            faiss.write_index(index, str(idx_path))
        with open(data_path, "w" if is_json else "wb", encoding="utf-8" if is_json else None) as f:
            if is_json:
                json.dump(data, f, ensure_ascii=False, indent=2)
            else:
                pickle.dump(data, f)

    @staticmethod
    def load(idx_path: Path, data_path: Path, is_json: bool = False) -> Tuple[Any, List]:
        """加载 FAISS 索引和对应的数据文件"""
        if idx_path.exists() and data_path.exists():
            try:
                index = faiss.read_index(str(idx_path))
                with open(data_path, "r" if is_json else "rb", encoding="utf-8" if is_json else None) as f:
                    data = json.load(f) if is_json else pickle.load(f)
                # 兼容性检测：旧索引可能是 IndexFlatL2 + 未归一化向量，需要丢弃并重建
                if isinstance(index, faiss.IndexFlatL2):
                    print(f"[WARN] 检测到旧版 L2 索引 ({idx_path})，将丢弃并触发重建")
                    return faiss.IndexFlatIP(1024), []
                return index, data
            except:
                pass
        return faiss.IndexFlatIP(1024), []

    # ---------- 微调数据管理 ----------
    @staticmethod
    def _read_existing_finetune_texts() -> set:
        """读取已有微调数据中的判例文本集合，用于去重"""
        existing = set()
        if PATHS.training_file.exists():
            try:
                with open(PATHS.training_file, "r", encoding="utf-8") as f:
                    for line in f:
                        try:
                            entry = json.loads(line)
                            user_msg = entry.get("messages", [{}])[1].get("content", "")
                            # 提取【待评分产品】和【参考标准】之间的文本
                            start = user_msg.find("【待评分产品】")
                            end = user_msg.find("【参考标准")
                            if start >= 0 and end > start:
                                text = user_msg[start + len("【待评分产品】"):end].strip()
                                if text:
                                    existing.add(text)
                        except:
                            continue
            except:
                pass
        return existing

    @staticmethod
    def append_cases_to_finetune(cases: List[Dict], sys_prompt: str, user_tpl: str) -> Tuple[int, int]:
        """
        将判例追加到微调数据文件（自动去重）。
        返回: (新增条数, 跳过的重复条数)
        """
        existing_texts = ResourceManager._read_existing_finetune_texts()
        added, skipped = 0, 0
        try:
            with open(PATHS.training_file, "a", encoding="utf-8") as f:
                for c in cases:
                    case_text = c.get("text", "").strip()
                    if case_text in existing_texts:
                        skipped += 1
                        continue
                    scores = c.get("scores", {})
                    master_comment = c.get("master_comment", "（人工校准）")
                    user_content = (user_tpl
                                    .replace("{product_desc}", case_text)
                                    .replace("{context_text}", "")
                                    .replace("{basic_case_text}", "")
                                    .replace("{case_text}", ""))
                    assistant_content = json.dumps(
                        {"master_comment": master_comment, "scores": scores},
                        ensure_ascii=False
                    )
                    entry = {
                        "messages": [
                            {"role": "system", "content": sys_prompt},
                            {"role": "user", "content": user_content},
                            {"role": "assistant", "content": assistant_content}
                        ]
                    }
                    f.write(json.dumps(entry, ensure_ascii=False) + "\n")
                    existing_texts.add(case_text)
                    added += 1
            return added, skipped
        except Exception as e:
            print(f"[ERROR] Finetune append failed: {e}")
            return 0, 0

    @staticmethod
    def save_ft_status(job_id, status, fine_tuned_model=None):
        data = {"job_id": job_id, "status": status, "timestamp": time.time()}
        if fine_tuned_model:
            data["fine_tuned_model"] = fine_tuned_model
        with open(PATHS.ft_status, 'w') as f:
            json.dump(data, f)

    @staticmethod
    def load_ft_status():
        if PATHS.ft_status.exists():
            try:
                return json.load(open(PATHS.ft_status, 'r'))
            except:
                pass
        return None

    # ---------- RAG 文件列表管理 ----------
    @staticmethod
    def save_kb_files(file_list: List[str]):
        with open(PATHS.kb_files, "w", encoding="utf-8") as f:
            json.dump(file_list, f, ensure_ascii=False, indent=2)

    @staticmethod
    def load_kb_files() -> List[str]:
        if PATHS.kb_files.exists():
            try:
                with open(PATHS.kb_files, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                pass
        return []


# ==========================================
# [SECTION 2] GitHub 同步工具 (GithubSync)
# ==========================================

class GithubSync:
    """负责将数据同步到 GitHub 仓库"""

    @staticmethod
    def _get_github_config():
        token = st.secrets.get("GITHUB_TOKEN")
        repo_name = st.secrets.get("GITHUB_REPO")
        branch = st.secrets.get("GITHUB_BRANCH", "main")
        return token, repo_name, branch

    @staticmethod
    def _get_github_client():
        token, repo_name, branch = GithubSync._get_github_config()
        if not token or not repo_name:
            return None, None, None
        g = Github(auth=Auth.Token(token))
        return g, repo_name, branch

    # ---------- 通用推送 / 删除 ----------
    @staticmethod
    def push_json(file_path_in_repo: str, data_dict, commit_msg: str = "Update via Streamlit") -> bool:
        g, repo_name, branch = GithubSync._get_github_client()
        if not g or not repo_name:
            st.error("❌ 未配置 Github Token 或 仓库名 (GITHUB_TOKEN / GITHUB_REPO)")
            return False
        try:
            repo = g.get_repo(repo_name)
            content_str = json.dumps(data_dict, ensure_ascii=False, indent=2)
            try:
                contents = repo.get_contents(file_path_in_repo, ref=branch)
                repo.update_file(path=contents.path, message=commit_msg, content=content_str, sha=contents.sha, branch=branch)
            except GithubException as e:
                if e.status == 404:
                    repo.create_file(path=file_path_in_repo, message=f"Create {file_path_in_repo}", content=content_str, branch=branch)
                else:
                    raise e
            return True
        except Exception as e:
            st.error(f"Github 同步失败: {str(e)}")
            return False

    @staticmethod
    def push_binary_file(file_path_in_repo: str, file_content: bytes, commit_msg: str = "Upload file") -> bool:
        g, repo_name, branch = GithubSync._get_github_client()
        if not g or not repo_name:
            st.error("❌ 未配置 Github Token 或 仓库名")
            return False
        try:
            repo = g.get_repo(repo_name)
            try:
                contents = repo.get_contents(file_path_in_repo, ref=branch)
                repo.update_file(path=contents.path, message=commit_msg, content=file_content, sha=contents.sha, branch=branch)
            except GithubException as e:
                if e.status == 404:
                    repo.create_file(path=file_path_in_repo, message=f"Create {file_path_in_repo}", content=file_content, branch=branch)
                else:
                    raise e
            return True
        except Exception as e:
            st.error(f"Github 文件上传失败: {str(e)}")
            return False

    @staticmethod
    def delete_file(file_path_in_repo: str, commit_msg: str = "Delete file") -> bool:
        g, repo_name, branch = GithubSync._get_github_client()
        if not g or not repo_name:
            return False
        try:
            repo = g.get_repo(repo_name)
            try:
                contents = repo.get_contents(file_path_in_repo, ref=branch)
                repo.delete_file(path=contents.path, message=commit_msg, sha=contents.sha, branch=branch)
                return True
            except GithubException as e:
                if e.status == 404:
                    return True
                raise e
        except Exception as e:
            st.error(f"Github 删除文件失败: {str(e)}")
            return False

    # ---------- 判例库同步（基础 / 进阶） ----------
    @staticmethod
    def sync_basic_cases(cases: List[Dict]) -> bool:
        return GithubSync.push_json("tea_data/basic_case.json", cases, "Update basic_case.json from App")

    @staticmethod
    def sync_supp_cases(cases: List[Dict]) -> bool:
        return GithubSync.push_json("tea_data/supplementary_case.json", cases, "Update supplementary_case.json from App")

    # ---------- RAG 文件管理 ----------
    @staticmethod
    def backup_rag_file(file_content: bytes, filename: str, backup_folder: str = "tea_backup") -> bool:
        file_path = f"{backup_folder}/{filename}"
        try:
            result = GithubSync.push_binary_file(file_path, file_content, f"Backup RAG file: {filename}")
            if result:
                print(f"[INFO] ✅ 已备份到 {file_path}")
            return result
        except Exception as e:
            print(f"[WARN] 备份文件 {filename} 到 {backup_folder} 失败: {e}")
            return False

    @staticmethod
    def add_rag_files(uploaded_files: List, rag_folder: str = "tea_data/RAG") -> Tuple[bool, List[str]]:
        g, repo_name, branch = GithubSync._get_github_client()
        if not g or not repo_name:
            st.error("❌ 未配置 Github Token 或 仓库名")
            return False, []
        try:
            uploaded_names = []
            for uf in uploaded_files:
                file_path = f"{rag_folder}/{uf.name}"
                uf.seek(0)
                file_content = uf.read()
                if GithubSync.push_binary_file(file_path, file_content, f"Add RAG file: {uf.name}"):
                    uploaded_names.append(uf.name)
                    GithubSync.backup_rag_file(file_content, uf.name, backup_folder="tea_backup")
                else:
                    st.warning(f"⚠️ 上传 {uf.name} 失败")
            return len(uploaded_names) > 0, uploaded_names
        except Exception as e:
            st.error(f"RAG文件添加失败: {str(e)}")
            return False, []

    @staticmethod
    def list_rag_files(rag_folder: str = "tea_data/RAG") -> List[str]:
        g, repo_name, branch = GithubSync._get_github_client()
        if not g or not repo_name:
            return []
        try:
            repo = g.get_repo(repo_name)
            contents = repo.get_contents(rag_folder, ref=branch)
            return [c.name for c in contents if c.type == "file"]
        except GithubException as e:
            if e.status == 404:
                return []
            print(f"[ERROR] 获取RAG文件列表失败: {e}")
            return []
        except Exception as e:
            print(f"[ERROR] 获取RAG文件列表失败: {e}")
            return []

    @staticmethod
    def delete_rag_file(filename: str, rag_folder: str = "tea_data/RAG") -> bool:
        file_path = f"{rag_folder}/{filename}"
        return GithubSync.delete_file(file_path, f"Delete RAG file: {filename}")

    @staticmethod
    def pull_rag_folder(rag_folder: str = "tea_data/RAG") -> List[Tuple[str, bytes]]:
        """从 GitHub 拉取 RAG 文件夹中的所有文件"""
        token, repo_name, branch = GithubSync._get_github_config()
        if not token or not repo_name:
            print("[WARN] GitHub config not found, skip pulling RAG")
            return []

        def download_with_retry(url, headers, max_retries=3):
            for attempt in range(1, max_retries + 1):
                try:
                    response = requests.get(url, headers=headers, timeout=180, stream=True)
                    if response.status_code == 200:
                        content = b''
                        for chunk in response.iter_content(chunk_size=8192):
                            if chunk:
                                content += chunk
                        return content, True
                    else:
                        print(f"[WARN]     尝试 {attempt}/{max_retries}: HTTP {response.status_code}")
                except Exception as e:
                    print(f"[WARN]     尝试 {attempt}/{max_retries}: {e}")
                    if attempt < max_retries:
                        time.sleep(2)
            return None, False

        try:
            g = Github(auth=Auth.Token(token))
            repo = g.get_repo(repo_name)
            files = []
            print(f"[INFO] ========== 开始从 GitHub 拉取 RAG 文件 ==========")

            try:
                contents = repo.get_contents(rag_folder, ref=branch)
                file_list = [c for c in contents if c.type == "file"]
                print(f"[INFO] 发现 {len(file_list)} 个文件")

                for idx, content in enumerate(file_list, 1):
                    print(f"\n[INFO] [{idx}/{len(file_list)}] 正在处理: {content.name}")
                    file_content = None
                    download_method = None

                    # 方法1：Raw URL
                    raw_url = f"https://raw.githubusercontent.com/{repo_name}/{branch}/{rag_folder}/{content.name}"
                    headers = {"Authorization": f"Bearer {token}"}
                    file_content, success = download_with_retry(raw_url, headers, max_retries=3)
                    if success and file_content:
                        download_method = "Raw URL"

                    # 方法2：Git Blob（小于1MB）
                    if file_content is None and content.size < 1024 * 1024:
                        try:
                            blob = repo.get_git_blob(content.sha)
                            if blob.encoding == "base64":
                                file_content = base64.b64decode(blob.content)
                                download_method = "Git Blob"
                        except Exception as e:
                            print(f"[WARN]   Git Blob 失败: {e}")

                    # 方法3：Download URL
                    if file_content is None and content.download_url:
                        headers = {"Authorization": f"Bearer {token}", "Accept": "application/vnd.github.v3.raw"}
                        file_content, success = download_with_retry(content.download_url, headers, max_retries=3)
                        if success:
                            download_method = "Download URL"

                    # 验证完整性
                    if file_content:
                        if len(file_content) == content.size:
                            files.append((content.name, file_content))
                            print(f"[INFO]   ✅ {content.name} 验证通过 ({download_method})")
                        else:
                            print(f"[ERROR]  ❌ {content.name} 大小不匹配")
                    else:
                        print(f"[ERROR]  ❌ {content.name} 所有下载方法均失败")

            except GithubException as e:
                if e.status == 404:
                    return []
                raise e

            print(f"\n[INFO] ========== RAG 拉取完成: {len(files)}/{len(file_list)} ==========\n")
            return files

        except Exception as e:
            print(f"[ERROR] 拉取 RAG 文件夹失败: {e}")
            import traceback
            traceback.print_exc()
            return []

    @staticmethod
    def download_github_file(file_path_in_repo: str) -> Optional[bytes]:
        """从 GitHub 下载单个文件的内容（用于下载模板等）"""
        token, repo_name, branch = GithubSync._get_github_config()
        if not token or not repo_name:
            return None
        try:
            g = Github(auth=Auth.Token(token))
            repo = g.get_repo(repo_name)
            content = repo.get_contents(file_path_in_repo, ref=branch)
            if content.encoding == "base64" and content.content:
                return base64.b64decode(content.content)
            # 大文件走 Raw URL
            raw_url = f"https://raw.githubusercontent.com/{repo_name}/{branch}/{file_path_in_repo}"
            resp = requests.get(raw_url, headers={"Authorization": f"Bearer {token}"}, timeout=30)
            if resp.status_code == 200:
                return resp.content
        except Exception as e:
            print(f"[WARN] 下载 {file_path_in_repo} 失败: {e}")
        return None


# ==========================================
# [SECTION 3] AI 服务 (Embedding & LLM)
# ==========================================

class AliyunEmbedder:
    """阿里云文本向量化服务"""
    def __init__(self, api_key):
        self.model_name = "text-embedding-v3"
        dashscope.api_key = api_key

    def encode(self, texts: List[str]) -> np.ndarray:
        if not texts:
            return np.zeros((0, 1024), dtype="float32")
        if isinstance(texts, str):
            texts = [texts]
        try:
            resp = TextEmbedding.call(model=self.model_name, input=texts)
            if resp.status_code == HTTPStatus.OK:
                embeddings = resp.output['embeddings']
                # 按 text_index 排序，保证向量顺序与输入文本顺序一致
                embeddings_sorted = sorted(embeddings, key=lambda x: x.get('text_index', 0))
                vecs = np.array([item['embedding'] for item in embeddings_sorted]).astype("float32")
                faiss.normalize_L2(vecs)  # 归一化，使内积等价于余弦相似度
                return vecs
            else:
                print(f"[ERROR] Embedding API 返回非200状态: code={resp.status_code}, msg={getattr(resp, 'message', 'unknown')}")
                st.warning(f"⚠️ Embedding 服务异常 (status={resp.status_code})，检索结果可能不准确")
        except Exception as e:
            print(f"[ERROR] Embedding API 调用失败: {e}")
            import traceback
            traceback.print_exc()
            st.warning(f"⚠️ Embedding 服务调用失败: {e}，检索结果可能不准确")
        return np.zeros((len(texts), 1024), dtype="float32")


def llm_normalize_user_input(raw_query: str, client: OpenAI) -> str:
    """使用 DeepSeek 对用户输入做语义清洗（去除非茶评内容）"""
    system_prompt = (
        """
          A. 角色与目标
          你是"茶评清洗器"。你的任务是从输入文本中提取并输出只与茶评相关的信息，删除无关内容，保持原意与原有表述风格，只能删减不能修改。
          B. 什么算"相关信息"（保留）
          仅保留与以下内容有关的句子/短语：
          茶的基本信息：茶名/品类、产地、年份、工艺、等级、原料、香型等
          干茶/茶汤/叶底：外观、色泽、条索、汤色、叶底描述
          香气与滋味：香气类型、强弱、层次、回甘、生津、涩感、苦感、甜度、醇厚度、喉韵、体感等
          冲泡信息与表现：器具、投茶量、水温、时间、出汤、几泡变化、耐泡度、适饮建议
          主观评价与结论：好喝/一般/缺点/性价比
          C. 什么算"无关信息"（删除）
          删除与茶评无直接关系的内容，例如：
          与茶无关的生活日常、情绪宣泄、社交聊天、段子
          店铺/物流/客服/包装破损/发货慢（除非"包装异味影响茶"这类直接影响品饮）
          广告、价格链接、优惠券、引流话术、品牌吹水（除非是"性价比"且与品饮结论相关）
          与其它产品/话题无关的对比闲聊
          凑字数内容
          D. 输出格式
          只输出清洗后的茶评正文，不要解释、不加标题、不输出"删除了什么"
          如果输入中没有任何茶评相关信息，则输出："无相关茶评信息"
          E. 操作原则
          尽量保留原句；只做删除/少量拼接
          不要补充不存在的细节，不要推测
          """
    )
    resp = client.chat.completions.create(
        model="deepseek-chat",
        temperature=0,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": raw_query}
        ]
    )
    return resp.choices[0].message.content.strip()


# ==========================================
# [SECTION 4] GraphRAG 静态知识库检索
# ==========================================

def _get_graphrag_artifact_dir() -> str:
    env_dir = os.getenv("GRAPHRAG_ARTIFACT_DIR", "").strip()
    if env_dir:
        return env_dir
    return str(getattr(PATHS, "GRAPHRAG_DIR", PATHS.DATA_DIR / "graphrag_artifacts"))


def _get_graphrag_retriever() -> 'GraphRAGRetriever | None':
    if st.session_state.get("_gr_retriever_loaded", False):
        return st.session_state.get("_gr_retriever_obj", None)
    artifact_dir = _get_graphrag_artifact_dir()
    edges_path = os.path.join(artifact_dir, "graph_edges.jsonl")
    comm_path = os.path.join(artifact_dir, "communities.json")
    if not (os.path.exists(edges_path) and os.path.exists(comm_path)):
        st.session_state["_gr_retriever_loaded"] = True
        st.session_state["_gr_retriever_obj"] = None
        return None
    try:
        gr = GraphRAGRetriever(artifact_dir=artifact_dir)
        st.session_state["_gr_retriever_loaded"] = True
        st.session_state["_gr_retriever_obj"] = gr
        return gr
    except Exception as e:
        print(f"[WARN] GraphRAGRetriever init failed: {e}")
        st.session_state["_gr_retriever_loaded"] = True
        st.session_state["_gr_retriever_obj"] = None
        return None


def graphrag_static_kb_context(query_vec: np.ndarray,
                               kb_index: faiss.Index,
                               kb_chunks: List[str],
                               k_num: Optional[int] = None,
                               top_seed: int = 5,
                               hop: int = 1,
                               max_expand: int = 12) -> Tuple[str, List[str]]:
    """使用向量种子 + GraphRAG 扩展构建知识库上下文"""
    if kb_index is None or getattr(kb_index, "ntotal", 0) <= 0 or not kb_chunks:
        return "（无手册资料）", []

    D, I = kb_index.search(query_vec, max(k_num, top_seed))
    vector_hits: List[Tuple[str, float]] = []
    for score, idx in zip(D[0].tolist(), I[0].tolist()):
        if idx is None or idx < 0 or idx >= len(kb_chunks):
            continue
        vector_hits.append((str(idx), float(score)))

    chunk_text_map = {str(i): kb_chunks[i] for i in range(len(kb_chunks))}

    gr = _get_graphrag_retriever()
    if gr is None:
        hits = [kb_chunks[int(cid)] for cid, _ in vector_hits[:k_num]]
        ctx = "\n".join([f"- {h[:240].strip()}..." for h in hits]) if hits else "（无手册资料）"
        return ctx, hits

    try:
        expanded = gr.expand(
            vector_hits=vector_hits,
            chunk_text_map=chunk_text_map,
            top_seed=top_seed, hop=hop, max_expand=max_expand,
            w_vec=0.7, w_graph=0.3
        )
        seeds = expanded.get("seed_chunks", [])
        exp_chunks = expanded.get("expanded_chunks", [])
        comm = expanded.get("community_summaries", [])
        seed_texts = [s.get("text", "") for s in seeds if s.get("text")]
        exp_texts = [c.get("text", "") for c in exp_chunks if c.get("text")]
        comm_texts = [c.get("summary", "") for c in comm if c.get("summary")]
        parts = []
        if comm_texts:
            parts.append("【GraphRAG 社区摘要】\n" + "\n\n".join(comm_texts[:2]))
        if seed_texts:
            parts.append("【向量检索种子片段】\n" + "\n".join([f"- {s[:240].strip()}..." for s in seed_texts[:k_num]]))
        if exp_texts:
            parts.append("【Graph 扩展片段】\n" + "\n".join([f"- {c[:240].strip()}..." for c in exp_texts[:k_num]]))
        ctx = "\n\n".join(parts) if parts else "（无手册资料）"
        hits_texts = []
        for t in (seed_texts + exp_texts):
            if t and t not in hits_texts:
                hits_texts.append(t)
            if len(hits_texts) >= k_num:
                break
        return ctx, hits_texts
    except Exception as e:
        print(f"[WARN] GraphRAG expand failed, fallback. err={e}")
        hits = [kb_chunks[int(cid)] for cid, _ in vector_hits[:k_num]]
        ctx = "\n".join([f"- {h[:240].strip()}..." for h in hits]) if hits else "（无手册资料）"
        return ctx, hits


# ==========================================
# [SECTION 5] 核心评分逻辑 (run_scoring)
# ==========================================

def _format_case_scores(sc: Dict) -> str:
    """将一条判例的所有因子评分格式化为文本"""
    parts = []
    for factor_name, factor_data in sc.items():
        if isinstance(factor_data, dict):
            score = factor_data.get('score', '?')
            comment = factor_data.get('comment', '')
            suggestion = factor_data.get('suggestion', '')
            parts.append(f"{factor_name}:{score}分 ({comment}; 建议:{suggestion})")
    return " | ".join(parts)


def run_scoring(text: str, kb_res: Tuple, basic_cases: List[Dict], supp_cases: Tuple,
                prompt_cfg: Dict, embedder: AliyunEmbedder, client: OpenAI,
                model_id: str, k_num: int, c_num: int):
    """
    执行 RAG 检索 + 判例组装 + LLM 评分。

    参数:
        text        : 待评分的茶评文本
        kb_res      : (kb_index, kb_chunks) 知识库
        basic_cases : 基础判例列表（全部加入 prompt）
        supp_cases  : (supp_index, supp_data) 进阶判例（做相似度匹配）
        prompt_cfg  : Prompt 配置 dict
        embedder    : 向量化服务
        client      : LLM 客户端
        model_id    : 模型标识
        k_num       : 参考知识库条目数
        c_num       : 参考进阶判例条目数

    返回: (scores_dict, kb_hits, found_supp_cases, system_prompt, user_prompt)
    """
    vec = embedder.encode([text])

    # 零向量检测：如果 embedding 失败，向量全为 0，检索结果将无意义
    if np.allclose(vec, 0):
        print(f"[WARN] 输入文本的 embedding 为零向量，检索结果将不可靠！text={text[:80]}...")
        st.warning("⚠️ 文本向量化失败（返回零向量），本次检索结果可能不准确，请检查阿里云 API Key 或网络连接。")

    # --- 知识库检索（GraphRAG） ---
    ctx_txt, hits = graphrag_static_kb_context(
        query_vec=vec,
        kb_index=kb_res[0],
        kb_chunks=kb_res[1],
        k_num=k_num,
        top_seed=max(5, k_num),
        hop=1,
        max_expand=12
    )

    # --- 基础判例：全部加入 prompt ---
    basic_case_text = "（无基础判例）"
    if basic_cases:
        parts = []
        for idx, bc in enumerate(basic_cases, 1):
            sc = bc.get('scores', {})
            score_str = _format_case_scores(sc)
            mc = bc.get('master_comment', '')
            parts.append(f"基础判例{idx}: {bc.get('text', '')[:300]}\n  评分: {score_str}\n  总评: {mc}")
        basic_case_text = "\n\n".join(parts)

    # --- 进阶判例：相似度匹配（仅比较 text 字段） ---
    case_txt = "（无相似进阶判例）"
    found_supp = []
    supp_index, supp_data = supp_cases
    if supp_index.ntotal > 0 and c_num > 0:
        _, idx_arr = supp_index.search(vec, min(c_num, supp_index.ntotal))
        supp_parts = []
        for i in idx_arr[0]:
            if 0 <= i < len(supp_data):
                c = supp_data[i]
                found_supp.append(c)
                sc = c.get('scores', {})
                score_str = _format_case_scores(sc)
                mc = c.get('master_comment', '')
                supp_parts.append(f"进阶判例: {c.get('text', '')[:300]}\n  评分: {score_str}\n  总评: {mc}")
        if supp_parts:
            case_txt = "\n\n".join(supp_parts)

    # --- 组装 Prompt ---
    sys_p = prompt_cfg.get('system_template', "")
    user_p = prompt_cfg.get('user_template', "")
    user_p = (user_p
              .replace("{product_desc}", text)
              .replace("{context_text}", ctx_txt)
              .replace("{basic_case_text}", basic_case_text)
              .replace("{case_text}", case_txt))

    # --- 调用 LLM ---
    try:
        resp = client.chat.completions.create(
            model=model_id,
            messages=[{"role": "system", "content": sys_p}, {"role": "user", "content": user_p}],
            response_format={"type": "json_object"},
            temperature=0.3
        )
        return json.loads(resp.choices[0].message.content), hits, found_supp, sys_p, user_p
    except Exception as e:
        st.error(f"Inference Error: {e}")
        return None, [], [], sys_p, user_p


# ==========================================
# [SECTION 6] 辅助工具与可视化
# ==========================================

def parse_file(uploaded_file) -> str:
    """解析上传文件为纯文本"""
    try:
        if uploaded_file.name.endswith('.txt'):
            return uploaded_file.read().decode("utf-8")
        if uploaded_file.name.endswith('.pdf'):
            return "".join([p.extract_text() for p in PdfReader(uploaded_file).pages])
        if uploaded_file.name.endswith('.docx'):
            return "\n".join([p.text for p in Document(uploaded_file).paragraphs])
    except:
        return ""
    return ""


def parse_file_bytes(filename: str, content: bytes) -> str:
    """解析文件内容（从 bytes）—— 用于从 GitHub 拉取的文件"""
    try:
        if filename.lower().endswith('.txt'):
            return content.decode('utf-8', errors='ignore')
        elif filename.lower().endswith('.pdf'):
            if not content.startswith(b'%PDF'):
                print(f"[ERROR] 不是有效的 PDF 文件: {filename}")
                return ""
            reader = PdfReader(BytesIO(content))
            text = ""
            for page in reader.pages:
                try:
                    pt = page.extract_text()
                    if pt:
                        text += pt + "\n"
                except:
                    continue
            return text
        elif filename.lower().endswith('.docx'):
            doc = Document(BytesIO(content))
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        print(f"[ERROR] 解析 {filename} 失败: {e}")
    return ""


def create_word_report(results: List[Dict]) -> BytesIO:
    """生成 Word 格式的批量评分报告"""
    doc = Document()
    doc.add_heading("茶评批量评分报告", 0)
    for item in results:
        doc.add_heading(f"条目 {item['id']}", 1)
        doc.add_paragraph(f"原文：{item['text']}")
        s = item.get('scores', {}).get('scores', {})
        mc = item.get('scores', {}).get('master_comment', '')
        if mc:
            doc.add_paragraph(f"总评：{mc}", style="Intense Quote")
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '因子', '分数', '评语', '建议'
        for k, v in s.items():
            r = table.add_row().cells
            r[0].text = k
            r[1].text = str(v.get('score', ''))
            r[2].text = v.get('comment', '')
            r[3].text = v.get('suggestion', '')
        doc.add_paragraph("_" * 20)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def plot_flavor_shape(scores_data: Dict):
    """绘制风味形态图（香调金字塔）"""
    s = scores_data["scores"]
    top = (s["优雅性"]["score"] + s["辨识度"]["score"]) / 2
    mid = (s["协调性"]["score"] + s["饱和度"]["score"]) / 2
    base = (s["持久性"]["score"] + s["苦涩度"]["score"]) / 2

    fig, ax = plt.subplots(figsize=(4, 5))
    fig.patch.set_alpha(0)
    ax.patch.set_alpha(0)

    y = np.array([1, 2, 3])
    x = np.array([base, mid, top])
    y_new = np.linspace(1, 3, 300)
    try:
        spl = make_interp_spline(y, x, k=2)
        x_smooth = spl(y_new)
    except:
        x_smooth = np.interp(y_new, y, x)
    x_smooth = np.maximum(x_smooth, 0.1)

    colors = {'base': '#8B4513', 'mid': '#D2691E', 'top': '#FFD700'}
    for mask, col in [((y_new >= 1.0) & (y_new <= 1.6), colors['base']),
                      ((y_new > 1.6) & (y_new <= 2.4), colors['mid']),
                      ((y_new > 2.4) & (y_new <= 3.0), colors['top'])]:
        ax.fill_betweenx(y_new[mask], -x_smooth[mask], x_smooth[mask], color=col, alpha=0.9, edgecolor=None)

    ax.plot(x_smooth, y_new, 'k', linewidth=1, alpha=0.2)
    ax.plot(-x_smooth, y_new, 'k', linewidth=1, alpha=0.2)
    ax.axhline(y=1.6, color='w', linestyle=':', alpha=0.5)
    ax.axhline(y=2.4, color='w', linestyle=':', alpha=0.5)

    font = {'ha': 'center', 'va': 'center', 'color': 'white', 'fontweight': 'bold', 'fontsize': 12}
    ax.text(0, 2.7, f"Top\n{top:.1f}", **font)
    ax.text(0, 2.0, f"Mid\n{mid:.1f}", **font)
    ax.text(0, 1.3, f"Base\n{base:.1f}", **font)
    ax.axis('off')
    ax.set_xlim(-10, 10)
    ax.set_ylim(0.8, 3.2)
    return fig


# ==========================================
# [SECTION 7] 判例初始化与 RAG 加载
# ==========================================

def bootstrap_cases(embedder: AliyunEmbedder):
    """
    初始化两类判例库（基础 + 进阶）。
    如果 session_state 为空，则从本地 JSON 文件加载。
    """
    # --- 基础判例 ---
    if len(st.session_state.basic_cases) == 0:
        seed = ResourceManager.load_external_json(PATHS.basic_case_data)
        if seed:
            st.session_state.basic_cases = seed
            ResourceManager.save_json(seed, PATHS.basic_case_data)

    # --- 进阶判例 ---
    supp_idx, supp_data = st.session_state.supp_cases
    if len(supp_data) == 0:
        seed = ResourceManager.load_external_json(PATHS.supp_case_data)
        if seed:
            texts = [c["text"] for c in seed]
            vecs = embedder.encode(texts)
            if supp_idx.ntotal == 0:
                supp_idx = faiss.IndexFlatIP(1024)
            if len(vecs) > 0:
                supp_idx.add(vecs)
            st.session_state.supp_cases = (supp_idx, seed)
            ResourceManager.save(supp_idx, seed, PATHS.supp_case_index, PATHS.supp_case_data, is_json=True)


def load_rag_from_github(aliyun_key: str) -> Tuple[bool, str]:
    """从 GitHub 拉取 RAG 文件并构建本地知识库"""
    print("\n[INFO] ========== 开始从 GitHub 加载 RAG 数据 ==========")
    try:
        rag_files = GithubSync.pull_rag_folder("tea_data/RAG")
        if not rag_files:
            return False, "GitHub 上没有找到 RAG 文件"

        all_text = ""
        file_names = []
        parse_failed = []
        for fname, fcontent in rag_files:
            file_names.append(fname)
            parsed = parse_file_bytes(fname, fcontent)
            if parsed and len(parsed.strip()) > 100:
                all_text += parsed + "\n"
            else:
                parse_failed.append(fname)

        if not all_text.strip():
            return False, "无法从 RAG 文件中提取有效文本"

        chunks = [all_text[i:i + 600] for i in range(0, len(all_text), 500)]
        if not chunks:
            return False, "切片失败"

        temp_embedder = AliyunEmbedder(aliyun_key)
        kb_idx = faiss.IndexFlatIP(1024)
        batch_size = 25
        all_vecs = []
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            try:
                vecs = temp_embedder.encode(batch)
                all_vecs.append(vecs)
            except Exception as e:
                print(f"[WARN] 批次向量化失败: {e}")

        if not all_vecs:
            return False, "向量化失败"

        vecs = np.vstack(all_vecs)
        kb_idx.add(vecs)

        st.session_state.kb = (kb_idx, chunks)
        st.session_state.kb_files = file_names
        ResourceManager.save(kb_idx, chunks, PATHS.kb_index, PATHS.kb_chunks)
        ResourceManager.save_kb_files(file_names)

        success_files = [f for f in file_names if f not in parse_failed]
        msg = f"✅ 成功加载 {len(chunks)} 条知识片段\n📁 来源: {', '.join(success_files)}"
        if parse_failed:
            msg += f"\n⚠️ 解析失败: {', '.join(parse_failed)}"
        return True, msg
    except Exception as e:
        return False, f"加载失败: {str(e)}"


# ==========================================
# [SECTION 8] 判例管理弹窗（基础 / 进阶通用）
# ==========================================

@st.dialog("📋 基础判例管理", width="large")
def show_basic_cases_dialog(embedder: AliyunEmbedder):
    """展示并管理基础判例"""
    cases = st.session_state.basic_cases
    if not cases:
        st.info("当前基础判例库为空")
        return

    st.write(f"共 **{len(cases)}** 条基础判例")
    st.caption("💡 勾选判例后，可以批量删除或转移到进阶判例库")

    selected = []
    for idx, case in enumerate(cases):
        with st.container(border=True):
            col1, col2 = st.columns([6, 1])
            with col1:
                text_preview = case.get('text', '')[:100] + ('...' if len(case.get('text', '')) > 100 else '')
                st.markdown(f"**#{idx + 1}** {text_preview}")
                scores = case.get('scores', {})
                if scores:
                    score_str = " | ".join([f"{k}:{v.get('score', '?')}" for k, v in scores.items()])
                    st.caption(score_str)
            with col2:
                if st.checkbox("选中", key=f"bc_sel_{idx}", label_visibility="collapsed"):
                    selected.append(idx)

    if selected:
        st.warning(f"已选中 {len(selected)} 条")
        act_c1, act_c2 = st.columns(2)
        with act_c1:
            if st.button("🗑️ 删除选中", type="primary", key="bc_del"):
                new_cases = [c for i, c in enumerate(cases) if i not in selected]
                st.session_state.basic_cases = new_cases
                ResourceManager.save_json(new_cases, PATHS.basic_case_data)
                with st.spinner("同步到GitHub..."):
                    GithubSync.sync_basic_cases(new_cases)
                st.success("删除完成！")
                time.sleep(1)
                st.rerun()
        with act_c2:
            if st.button("➡️ 转移到进阶判例", key="bc_to_supp"):
                moved = [cases[i] for i in selected]
                new_basic = [c for i, c in enumerate(cases) if i not in selected]
                # 更新基础判例
                st.session_state.basic_cases = new_basic
                ResourceManager.save_json(new_basic, PATHS.basic_case_data)
                # 添加到进阶判例
                supp_idx, supp_data = st.session_state.supp_cases
                supp_data.extend(moved)
                moved_vecs = embedder.encode([c["text"] for c in moved])
                supp_idx.add(moved_vecs)
                st.session_state.supp_cases = (supp_idx, supp_data)
                ResourceManager.save(supp_idx, supp_data, PATHS.supp_case_index, PATHS.supp_case_data, is_json=True)
                with st.spinner("同步到GitHub..."):
                    GithubSync.sync_basic_cases(new_basic)
                    GithubSync.sync_supp_cases(supp_data)
                st.success(f"已转移 {len(moved)} 条到进阶判例！")
                time.sleep(1)
                st.rerun()


@st.dialog("📋 进阶判例管理", width="large")
def show_supp_cases_dialog(embedder: AliyunEmbedder):
    """展示并管理进阶判例"""
    _, cases = st.session_state.supp_cases
    if not cases:
        st.info("当前进阶判例库为空")
        return

    st.write(f"共 **{len(cases)}** 条进阶判例")
    st.caption("💡 勾选判例后，可以批量删除或转移到基础判例库")

    selected = []
    for idx, case in enumerate(cases):
        with st.container(border=True):
            col1, col2 = st.columns([6, 1])
            with col1:
                text_preview = case.get('text', '')[:100] + ('...' if len(case.get('text', '')) > 100 else '')
                st.markdown(f"**#{idx + 1}** {text_preview}")
                scores = case.get('scores', {})
                if scores:
                    score_str = " | ".join([f"{k}:{v.get('score', '?')}" for k, v in scores.items()])
                    st.caption(score_str)
            with col2:
                if st.checkbox("选中", key=f"sc_sel_{idx}", label_visibility="collapsed"):
                    selected.append(idx)

    if selected:
        st.warning(f"已选中 {len(selected)} 条")
        act_c1, act_c2 = st.columns(2)
        with act_c1:
            if st.button("🗑️ 删除选中", type="primary", key="sc_del"):
                new_data = [c for i, c in enumerate(cases) if i not in selected]
                new_idx = faiss.IndexFlatIP(1024)
                if new_data:
                    vecs = embedder.encode([c["text"] for c in new_data])
                    new_idx.add(vecs)
                st.session_state.supp_cases = (new_idx, new_data)
                ResourceManager.save(new_idx, new_data, PATHS.supp_case_index, PATHS.supp_case_data, is_json=True)
                with st.spinner("同步到GitHub..."):
                    GithubSync.sync_supp_cases(new_data)
                st.success("删除完成！")
                time.sleep(1)
                st.rerun()
        with act_c2:
            if st.button("⬅️ 转移到基础判例", key="sc_to_basic"):
                moved = [cases[i] for i in selected]
                new_supp = [c for i, c in enumerate(cases) if i not in selected]
                # 重建进阶索引
                new_idx = faiss.IndexFlatIP(1024)
                if new_supp:
                    vecs = embedder.encode([c["text"] for c in new_supp])
                    new_idx.add(vecs)
                st.session_state.supp_cases = (new_idx, new_supp)
                ResourceManager.save(new_idx, new_supp, PATHS.supp_case_index, PATHS.supp_case_data, is_json=True)
                # 添加到基础判例
                st.session_state.basic_cases.extend(moved)
                ResourceManager.save_json(st.session_state.basic_cases, PATHS.basic_case_data)
                with st.spinner("同步到GitHub..."):
                    GithubSync.sync_supp_cases(new_supp)
                    GithubSync.sync_basic_cases(st.session_state.basic_cases)
                st.success(f"已转移 {len(moved)} 条到基础判例！")
                time.sleep(1)
                st.rerun()


@st.dialog("✏️ 编辑基础判例", width="large")
def edit_basic_case_dialog(case_idx: int):
    """编辑单个基础判例"""
    cases = st.session_state.basic_cases
    if case_idx >= len(cases):
        st.error("判例不存在")
        return
    case = cases[case_idx]
    factors = ["优雅性", "辨识度", "协调性", "饱和度", "持久性", "苦涩度"]
    st.subheader(f"编辑基础判例 #{case_idx + 1}")
    new_text = st.text_area("判例描述", case.get("text", ""), height=100)
    new_master = st.text_area("宗师总评", case.get("master_comment", ""), height=60)
    st.markdown("**因子评分**")
    new_scores = {}
    cols = st.columns(3)
    old_scores = case.get("scores", {})
    for i, f in enumerate(factors):
        with cols[i % 3]:
            with st.container(border=True):
                st.markdown(f"**{f}**")
                old_f = old_scores.get(f, {})
                new_scores[f] = {
                    "score": st.number_input("分数", 0, 9, int(old_f.get("score", 5)), key=f"ebc_s_{f}"),
                    "comment": st.text_input("评语", old_f.get("comment", ""), key=f"ebc_c_{f}"),
                    "suggestion": st.text_input("建议", old_f.get("suggestion", ""), key=f"ebc_sg_{f}")
                }
    if st.button("💾 保存修改并同步", type="primary"):
        cases[case_idx] = {
            "text": new_text, "scores": new_scores,
            "master_comment": new_master, "created_at": case.get("created_at", time.strftime("%Y-%m-%d"))
        }
        st.session_state.basic_cases = cases
        ResourceManager.save_json(cases, PATHS.basic_case_data)
        with st.spinner("同步到GitHub..."):
            GithubSync.sync_basic_cases(cases)
        st.session_state.editing_basic_idx = None
        st.success("保存成功！")
        time.sleep(1)
        st.rerun()


@st.dialog("✏️ 编辑进阶判例", width="large")
def edit_supp_case_dialog(case_idx: int, embedder: AliyunEmbedder):
    """编辑单个进阶判例"""
    _, cases = st.session_state.supp_cases
    if case_idx >= len(cases):
        st.error("判例不存在")
        return
    case = cases[case_idx]
    factors = ["优雅性", "辨识度", "协调性", "饱和度", "持久性", "苦涩度"]
    st.subheader(f"编辑进阶判例 #{case_idx + 1}")
    new_text = st.text_area("判例描述", case.get("text", ""), height=100)
    new_master = st.text_area("宗师总评", case.get("master_comment", ""), height=60)
    st.markdown("**因子评分**")
    new_scores = {}
    cols = st.columns(3)
    old_scores = case.get("scores", {})
    for i, f in enumerate(factors):
        with cols[i % 3]:
            with st.container(border=True):
                st.markdown(f"**{f}**")
                old_f = old_scores.get(f, {})
                new_scores[f] = {
                    "score": st.number_input("分数", 0, 9, int(old_f.get("score", 5)), key=f"esc_s_{f}"),
                    "comment": st.text_input("评语", old_f.get("comment", ""), key=f"esc_c_{f}"),
                    "suggestion": st.text_input("建议", old_f.get("suggestion", ""), key=f"esc_sg_{f}")
                }
    if st.button("💾 保存修改并同步", type="primary"):
        cases[case_idx] = {
            "text": new_text, "scores": new_scores,
            "master_comment": new_master, "created_at": case.get("created_at", time.strftime("%Y-%m-%d"))
        }
        # 重建 FAISS 索引
        new_idx = faiss.IndexFlatIP(1024)
        texts = [c["text"] for c in cases]
        vecs = embedder.encode(texts)
        new_idx.add(vecs)
        st.session_state.supp_cases = (new_idx, cases)
        ResourceManager.save(new_idx, cases, PATHS.supp_case_index, PATHS.supp_case_data, is_json=True)
        with st.spinner("同步到GitHub..."):
            GithubSync.sync_supp_cases(cases)
        st.session_state.editing_supp_idx = None
        st.success("保存成功！")
        time.sleep(1)
        st.rerun()


# ==========================================
# [SECTION 9] 提示词查看弹窗 & 茶评示例弹窗
# ==========================================

@st.dialog("📝 本次发送给LLM的完整Prompt", width="large")
def show_prompt_dialog():
    """弹窗展示发送给 LLM 的系统提示词和用户提示词"""
    st.markdown("**🔧 System Prompt（系统提示词）：**")
    st.code(st.session_state.get('last_llm_sys_prompt', '（暂无）'), language=None)
    st.markdown("**💬 User Prompt（用户提示词）：**")
    st.code(st.session_state.get('last_llm_user_prompt', '（暂无）'), language=None)


TEA_EXAMPLES = [
    {
        "title": "🌸 桂花乌龙",
        "text": "这款桂花乌龙干茶清甜带花香，热闻像刚蒸好的桂花糕。入口先是乌龙的清爽与微焙火，随后桂花香在回甘里慢慢铺开；茶汤金黄透亮，喉韵干净不腻。适合下午提神，也很适合配低糖点心。"
    },
    {
        "title": "🔥 正山小种红茶",
        "text": "这款正山小种红茶干茶散发淡淡的松烟香与蜜甜香；冲泡后汤色橙红明亮，通透如琥珀，金圈明显。香气初闻是浓郁的桂圆干香，伴随传统的松烟熏香，稍凉后溢出野花蜜甜。入口醇厚顺滑，烟熏味与甜润感平衡得宜，喉韵甘甜持久，回甘中带着隐约的果脯香。叶底呈古铜色，叶片柔韧有光泽，匀整度高，整体传统工艺到位，烟熏感不呛喉，甜感突出，适合喜好醇厚口感的茶客。"
    },
    {
        "title": "⚠️ 茉莉花茶（反面案例）",
        "text": "这款茉莉花茶干茶中花瓣残留过多，花香浮于表面不自然；汤色浅黄泛绿，稍显浑浊。香气初闻茉莉香刺鼻，实则香精感明显，冲泡两三次后香气断崖式下跌。滋味上茶汤寡淡，花香与茶味分离，涩感明显，饮后口腔有粘腻感。叶底花渣混杂，茶叶粗硬，完整性差。茶底工艺粗糙，缺乏传统窨制应有的'冰糖甜'。"
    },
    {
        "title": "🤍 白牡丹",
        "text": "这款白牡丹一冲开就带着淡淡的花香和干净的毫香，茶汤清亮柔和，入口像温润的梨子水般清甜，细细咽下去喉咙里有轻微的凉意，回甘不急不躁却很持久，越往后泡越显出淡雅的草木气。"
    },
    {
        "title": "🍂 普洱熟茶",
        "text": "这杯普洱熟茶汤色红浓透亮，陈香里夹着糯甜和一点点木质气息，入口醇厚顺滑几乎没有涩感，暖意从胃里慢慢散开，尾韵干净，杯底带着淡淡的枣香，适合在忙碌的晚上安安稳稳地喝完一壶。"
    }
]


@st.dialog("🍵 茶评示例", width="large")
def show_tea_examples_dialog():
    """展示预置茶评示例文本"""
    st.caption("以下是五组茶评示例，点击文本框即可选中复制，粘贴到「交互评分」中使用。")
    for i, ex in enumerate(TEA_EXAMPLES):
        st.markdown(f"**{ex['title']}**")
        st.code(ex["text"], language=None)
        if i < len(TEA_EXAMPLES) - 1:
            st.markdown("")


# ==========================================
# [SECTION 10] 模板下载辅助函数
# ==========================================

def get_template_bytes() -> Optional[bytes]:
    """获取模板文件内容（优先本地，其次从 GitHub 下载）"""
    if PATHS.template_file.exists():
        with open(PATHS.template_file, 'rb') as f:
            return f.read()
    # 尝试从 GitHub 下载
    content = GithubSync.download_github_file("tea_backup/template.xlsx")
    if content:
        # 缓存到本地
        try:
            with open(PATHS.template_file, 'wb') as f:
                f.write(content)
        except:
            pass
    return content


# ==========================================
# [SECTION 11] 主程序 —— Session 初始化
# ==========================================

if 'loaded' not in st.session_state:
    print("\n" + "=" * 70)
    print("[INFO] ========== 茶饮六因子AI评分器 - 系统初始化 ==========")
    print("=" * 70)

    # 1. 加载知识库缓存
    print("[INFO] 步骤 1/4: 加载知识库缓存...")
    kb_idx, kb_data = ResourceManager.load(PATHS.kb_index, PATHS.kb_chunks)
    st.session_state.kb = (kb_idx, kb_data)
    st.session_state.kb_files = ResourceManager.load_kb_files()
    print(f"[INFO]   → 知识库: {len(kb_data)} 个片段")

    # 2. 加载判例库（基础 + 进阶）
    print("[INFO] 步骤 2/4: 加载判例库...")
    st.session_state.basic_cases = ResourceManager.load_external_json(PATHS.basic_case_data, fallback=[])
    supp_idx, supp_data = ResourceManager.load(PATHS.supp_case_index, PATHS.supp_case_data, is_json=True)
    st.session_state.supp_cases = (supp_idx, supp_data)
    print(f"[INFO]   → 基础判例: {len(st.session_state.basic_cases)} 条")
    print(f"[INFO]   → 进阶判例: {len(supp_data)} 条")

    # 3. RAG 延迟加载标记
    print("[INFO] 步骤 3/4: 检查 RAG 状态...")
    if len(kb_data) == 0:
        st.session_state.rag_loading_needed = True
        st.session_state.rag_loading_status = "pending"
        print("[INFO]   ⚠️ 本地知识库为空，将从 GitHub 加载")
    else:
        st.session_state.rag_loading_needed = False
        st.session_state.rag_loading_status = "complete"
        print(f"[INFO]   ✅ 使用本地缓存: {len(kb_data)} 个片段")

    # 4. Prompt 配置
    print("[INFO] 步骤 4/4: 加载 Prompt 配置...")
    if PATHS.prompt_config_file.exists():
        try:
            with open(PATHS.prompt_config_file, 'r', encoding='utf-8') as f:
                st.session_state.prompt_config = json.load(f)
                print("[INFO]   ✅ 已加载自定义 Prompt 配置")
        except Exception as e:
            print(f"[WARN]   加载失败: {e}")

    if 'prompt_config' not in st.session_state:
        sys_prompt_content = ResourceManager.load_external_text(PATHS.SRC_SYS_PROMPT, fallback="你是一名茶评专家...")
        st.session_state.prompt_config = {
            "system_template": sys_prompt_content,
            "user_template": DEFAULT_USER_TEMPLATE
        }
        print("[INFO]   ✅ 使用默认 Prompt 配置")

    st.session_state.loaded = True
    print("=" * 70)
    print("[INFO] ========== 系统初始化完成 ==========")
    print("=" * 70 + "\n")


# ==========================================
# [SECTION 12] 侧边栏
# ==========================================

with st.sidebar:
    st.header("⚙️ 系统配置")

    # --- API 配置 ---
    st.markdown("**🔐 API 配置**")
    aliyun_key = os.getenv("ALIYUN_API_KEY") or st.secrets.get("ALIYUN_API_KEY", "")
    deepseek_key = os.getenv("DEEPSEEK_API_KEY") or st.secrets.get("DEEPSEEK_API_KEY", "")

    if not aliyun_key or not deepseek_key:
        st.warning("⚠️ 未配置 API Key")
        st.stop()
    else:
        st.success("✅ API 就绪")

    st.markdown("---")

    # --- 模型状态 ---
    st.markdown(f"**预处理模型：** `Deepseek-chat`")
    st.markdown(f"**评分模型：** `Qwen3-14B`")
    model_id = "Qwen3-14B"
    try:
        resp = requests.get("http://117.50.138.123:8001/status", timeout=2)
        if resp.status_code == 200 and resp.json().get("lora_available"):
            model_id = "default_lora"
            st.success("🎉 已启用微调模型")
    except:
        pass
    ft_status = ResourceManager.load_ft_status()
    if ft_status and ft_status.get("status") == "succeeded":
        st.info(f"🎉 发现微调模型：`{ft_status.get('fine_tuned_model')}`")

    # --- 初始化服务实例 ---
    embedder = AliyunEmbedder(aliyun_key)
    client = OpenAI(api_key="dummy", base_url="http://117.50.138.123:8000/v1")
    client_d = OpenAI(api_key=deepseek_key, base_url="https://api.deepseek.com")

    bootstrap_cases(embedder)

    st.markdown("---")

    # --- 延迟加载 RAG ---
    kb_count = len(st.session_state.kb[1])
    basic_count = len(st.session_state.basic_cases)
    supp_count = len(st.session_state.supp_cases[1])

    if st.session_state.get('rag_loading_needed', False):
        loading_status = st.session_state.get('rag_loading_status', 'pending')
        if loading_status == 'pending':
            with st.status("🔄 正在从 GitHub 加载知识库...", expanded=True) as status:
                st.write("📥 下载 RAG 文件...")
                st.session_state.rag_loading_status = 'loading'
                try:
                    success, msg = load_rag_from_github(aliyun_key)
                    if success:
                        status.update(label="✅ 知识库加载完成", state="complete", expanded=False)
                        st.session_state.rag_loading_status = 'complete'
                        st.session_state.rag_loading_needed = False
                        time.sleep(1)
                        st.rerun()
                    else:
                        status.update(label="❌ 知识库加载失败", state="error", expanded=True)
                        st.error(msg)
                        st.info("💡 您可以在「知识库设计」手动上传 RAG 文件")
                        st.session_state.rag_loading_status = 'failed'
                        if st.button("🔄 重试加载", type="secondary"):
                            st.session_state.rag_loading_status = 'pending'
                            st.rerun()
                except Exception as e:
                    status.update(label="❌ 加载出错", state="error", expanded=True)
                    st.error(f"加载失败: {str(e)}")
                    st.session_state.rag_loading_status = 'failed'
                    if st.button("🔄 重试加载", type="secondary"):
                        st.session_state.rag_loading_status = 'pending'
                        st.rerun()
        elif loading_status == 'loading':
            st.info("🔄 正在加载知识库，请稍候...")
        elif loading_status == 'failed':
            st.warning("⚠️ 知识库加载失败")
            if st.button("🔄 重试从 GitHub 加载", type="secondary"):
                st.session_state.rag_loading_status = 'pending'
                st.rerun()

    kb_count = len(st.session_state.kb[1])
    kb_files = st.session_state.get('kb_files', [])

    st.markdown(f"知识库：**{kb_count}** 个（向量片段）")
    st.markdown(f"基础判例：**{basic_count}** 条")
    st.markdown(f"进阶判例：**{supp_count}** 条")
    if kb_count == 0 and not kb_files:
        st.caption("⚠️ 知识库为空，请上传文件或从云端加载")

    st.markdown("---")
    if st.button("🍵 茶评示例", use_container_width=True):
        show_tea_examples_dialog()


# ==========================================
# [SECTION 13] 主界面 —— 标题 & Tab 定义
# ==========================================

st.markdown('<div class="main-title">🍵 茶品六因子 AI 评分器 Pro</div>', unsafe_allow_html=True)
st.markdown('<div class="slogan">"一片叶子落入水中，改变了水的味道..."</div>', unsafe_allow_html=True)
st.markdown('<p style="text-align:center; color:#888; font-size:0.95em;">推理服务开放时间：9:00~20:00</p>', unsafe_allow_html=True)

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "💡 交互评分", "🚀 批量评分", "📕 知识库设计",
    "📋 判例库设计", "🔧 模型微调", "📲 提示词（Prompt）配置"
])


# ==========================================
# [TAB 1] 交互评分
# ==========================================

with tab1:
    st.info("将参考知识库与判例库进行评分。确认结果可一键更新判例库。")
    c1, c2, c3, c4, c5 = st.columns([1, 3, 1, 3, 1])
    r_num = c2.number_input("参考知识库条目数量", 1, 20, 3, key="r1")
    c_num = c4.number_input("参考进阶判例条目数量", 1, 20, 5, key="c1")

    if 'current_user_input' not in st.session_state:
        st.session_state.current_user_input = ""
    user_input = st.text_area("请输入茶评描述:", value=st.session_state.current_user_input, height=150, key="ui")
    st.session_state.current_user_input = user_input

    if 'last_scores' not in st.session_state:
        st.session_state.last_scores = None
        st.session_state.last_master_comment = ""
    if 'last_llm_sys_prompt' not in st.session_state:
        st.session_state.last_llm_sys_prompt = ""
    if 'last_llm_user_prompt' not in st.session_state:
        st.session_state.last_llm_user_prompt = ""
    if 'score_version' not in st.session_state:
        st.session_state.score_version = 0

    if st.button("开始评分", type="primary", use_container_width=True):
        if not user_input:
            st.warning("请输入内容")
        else:
            with st.spinner(f"正在使用 {model_id} 品鉴..."):
                user_input_clean = llm_normalize_user_input(user_input, client_d)
                scores, kb_h, case_h, sent_sys_p, sent_user_p = run_scoring(
                    user_input_clean, st.session_state.kb,
                    st.session_state.basic_cases, st.session_state.supp_cases,
                    st.session_state.prompt_config, embedder, client, "Qwen3-14B", r_num, c_num
                )
                # 无论是否成功，都保存 prompt
                st.session_state.last_llm_sys_prompt = sent_sys_p
                st.session_state.last_llm_user_prompt = sent_user_p
                if scores:
                    st.session_state.last_scores = scores
                    st.session_state.last_master_comment = scores.get("master_comment", "")
                    st.session_state.score_version += 1
                else:
                    st.session_state.last_scores = None
                    st.session_state.last_master_comment = ""
                st.rerun()

    # --- 提示词查看弹窗（在评分结果之上） ---
    if st.session_state.last_llm_sys_prompt or st.session_state.last_llm_user_prompt:
        if st.button("📝 查看本次发送给LLM的提示词"):
            show_prompt_dialog()

    # --- 评分结果展示 ---
    if st.session_state.last_scores:
        s = st.session_state.last_scores["scores"]
        mc = st.session_state.last_master_comment
        st.markdown(f'<div class="master-comment"><b>👵 宗师总评：</b><br>{mc}</div>', unsafe_allow_html=True)

        left_col, right_col = st.columns([35, 65])
        with left_col:
            st.subheader("📊 风味形态")
            st.pyplot(plot_flavor_shape(st.session_state.last_scores), use_container_width=True)
        with right_col:
            cols = st.columns(2)
            factors = ["优雅性", "辨识度", "协调性", "饱和度", "持久性", "苦涩度"]
            for i, f in enumerate(factors):
                if f in s:
                    d = s[f]
                    with cols[i % 2]:
                        st.markdown(f"""<div class="factor-card"><div class="score-header"><span>{f}</span><span>{d['score']}/9</span></div><div>{d['comment']}</div><div class="advice-tag">💡 {d.get('suggestion', '')}</div></div>""", unsafe_allow_html=True)

        # --- 校准与修正 ---
        st.subheader("🛠️ 评分校准与修正")
        v = st.session_state.score_version
        cal_master = st.text_area("校准总评", mc, key=f"cal_master_{v}")
        cal_scores = {}
        st.write("分项调整")
        active_factors = [f for f in factors if f in s]
        grid_cols = st.columns(3)
        for i, f in enumerate(active_factors):
            with grid_cols[i % 3]:
                with st.container(border=True):
                    t_col, s_col = st.columns([1, 1])
                    with t_col:
                        st.markdown(f"<div style='padding-top: 5px;'><b>📌 {f}</b></div>", unsafe_allow_html=True)
                    with s_col:
                        new_score = st.number_input("分数", 0, 9, int(s[f]['score']), 1, key=f"s_{f}_{v}", label_visibility="collapsed")
                    cal_scores[f] = {
                        "score": new_score,
                        "comment": st.text_area("评语", s[f]['comment'], key=f"c_{f}_{v}", height=80, placeholder="评语"),
                        "suggestion": st.text_area("建议", s[f].get('suggestion', ''), key=f"sg_{f}_{v}", height=68, placeholder="建议")
                    }

        if st.button("💾 保存校准评分", type="primary"):
            nc = {
                "text": user_input, "scores": cal_scores,
                "master_comment": cal_master, "created_at": time.strftime("%Y-%m-%d")
            }
            # 校准结果默认保存到进阶判例
            supp_idx, supp_data = st.session_state.supp_cases
            supp_data.append(nc)
            supp_idx.add(embedder.encode([user_input]))
            st.session_state.supp_cases = (supp_idx, supp_data)
            ResourceManager.save(supp_idx, supp_data, PATHS.supp_case_index, PATHS.supp_case_data, is_json=True)
            with st.spinner("同步判例到GitHub..."):
                GithubSync.sync_supp_cases(supp_data)
            st.success("校准已保存到进阶判例并同步")
            st.rerun()
    elif st.session_state.last_llm_sys_prompt and st.session_state.last_scores is None:
        # prompt 已发送但模型未返回有效结果
        st.warning("⚠️ 模型未返回有效评分结果，请检查上方的提示词内容或重试。")


# ==========================================
# [TAB 2] 批量评分
# ==========================================

with tab2:
    c1, c2, c3, c4, c5 = st.columns([1, 3, 1, 3, 1])
    r_n = c2.number_input("参考知识库条目数量", 1, 20, 3, key="rb")
    c_n = c4.number_input("参考进阶判例条目数量", 1, 20, 5, key="cb")

    f = st.file_uploader("上传文件 (.txt/.docx)", key="batch_uploader")

    if f and st.button("批量处理"):
        lines = [l.strip() for l in parse_file(f).split('\n') if len(l) > 10]
        res, bar = [], st.progress(0)
        for i, l in enumerate(lines):
            l = llm_normalize_user_input(l, client_d)
            s, _, _, _, _ = run_scoring(
                l, st.session_state.kb,
                st.session_state.basic_cases, st.session_state.supp_cases,
                st.session_state.prompt_config, embedder, client, "Qwen3-14B", r_n, c_n
            )
            res.append({"id": i + 1, "text": l, "scores": s})
            bar.progress((i + 1) / len(lines))
        st.success("完成")
        st.download_button("下载Word", create_word_report(res), "report.docx")


# ==========================================
# [TAB 3] 知识库设计（RAG）—— 不做改动
# ==========================================

with tab3:
    st.subheader("📚 知识库 (RAG)")
    st.caption("上传PDF/文档以增强模型回答的准确性。文件将同步到云端。添加或删除文件后，系统将自动从云端重建本地知识库。")

    # ===== GitHub RAG 文件列表 =====
    st.markdown("**📁 云端上的RAG文件：**")
    if 'github_rag_files' not in st.session_state:
        st.session_state.github_rag_files = []

    col_refresh, col_spacer = st.columns([1, 3])
    with col_refresh:
        if st.button("🔄 刷新列表", key="refresh_rag_list"):
            with st.spinner("正在获取文件列表..."):
                st.session_state.github_rag_files = GithubSync.list_rag_files()
            st.rerun()

    github_files = st.session_state.github_rag_files
    if not github_files:
        github_files = GithubSync.list_rag_files()
        st.session_state.github_rag_files = github_files

    if github_files:
        st.info(f"共 {len(github_files)} 个文件")
        if 'rag_files_to_delete' not in st.session_state:
            st.session_state.rag_files_to_delete = set()

        for fname in github_files:
            file_col, del_col = st.columns([5, 1])
            with file_col:
                if fname in st.session_state.rag_files_to_delete:
                    st.markdown(f"~~📄 {fname}~~ *(待删除)*")
                else:
                    st.markdown(f"📄 {fname}")
            with del_col:
                if fname not in st.session_state.rag_files_to_delete:
                    if st.button("🗑️", key=f"del_rag_{fname}", help=f"删除 {fname}"):
                        st.session_state.rag_files_to_delete.add(fname)
                        st.rerun()
                else:
                    if st.button("↩️", key=f"undo_rag_{fname}", help="撤销删除"):
                        st.session_state.rag_files_to_delete.discard(fname)
                        st.rerun()

        if st.session_state.rag_files_to_delete:
            st.warning(f"⚠️ 将删除 {len(st.session_state.rag_files_to_delete)} 个文件")
            del_col1, del_col2 = st.columns(2)
            with del_col1:
                if st.button("✅ 确认删除并同步知识库", type="primary", key="confirm_del_rag"):
                    with st.spinner("正在删除文件..."):
                        deleted = []
                        for fname in st.session_state.rag_files_to_delete:
                            if GithubSync.delete_rag_file(fname):
                                deleted.append(fname)
                        st.session_state.github_rag_files = [f for f in github_files if f not in deleted]
                        current_kb_files = st.session_state.get('kb_files', [])
                        st.session_state.kb_files = [f for f in current_kb_files if f not in deleted]
                        ResourceManager.save_kb_files(st.session_state.kb_files)
                        st.session_state.rag_files_to_delete = set()
                        st.success(f"✅ 已删除 {len(deleted)} 个文件")
                    with st.spinner("🔄 正在从云端重建本地知识库..."):
                        success, msg = load_rag_from_github(aliyun_key)
                        if success:
                            st.success(msg)
                            st.session_state.github_rag_files = GithubSync.list_rag_files()
                        else:
                            st.warning(f"知识库重建失败: {msg}")
                    time.sleep(1)
                    st.rerun()
            with del_col2:
                if st.button("❌ 取消", key="cancel_del_rag"):
                    st.session_state.rag_files_to_delete = set()
                    st.rerun()
    else:
        st.caption("暂无RAG文件")

    st.markdown("---")

    # ===== 上传新文件 =====
    st.markdown("**➕ 添加新文件：**")
    up = st.file_uploader("选择文件", accept_multiple_files=True, key="kb_uploader", type=['pdf', 'txt', 'docx'])

    if up and st.button("📤 添加到知识库并同步", type="primary"):
        new_names = [u.name for u in up]
        existing_names = st.session_state.get('github_rag_files', [])
        duplicate_names = set(new_names) & set(existing_names)
        if duplicate_names:
            st.warning(f"⚠️ 以下文件已存在，将被覆盖：{', '.join(duplicate_names)}")
        with st.spinner("正在处理文件..."):
            raw = "".join([parse_file(u) for u in up])
            if not raw.strip():
                st.error("❌ 无法从上传的文件中提取有效文本")
            else:
                with st.spinner("上传到GitHub..."):
                    success, uploaded_names = GithubSync.add_rag_files(up, "tea_data/RAG")
                if success:
                    current_kb_files = st.session_state.get('kb_files', [])
                    all_files = list(set(current_kb_files + uploaded_names))
                    st.session_state.kb_files = all_files
                    st.session_state.github_rag_files = list(set(existing_names + uploaded_names))
                    ResourceManager.save_kb_files(all_files)
                    st.success(f"✅ 已上传 {len(uploaded_names)} 个文件到GitHub")
                    with st.spinner("🔄 正在从云端重建本地知识库..."):
                        rebuild_success, rebuild_msg = load_rag_from_github(aliyun_key)
                        if rebuild_success:
                            st.success(rebuild_msg)
                            st.session_state.github_rag_files = GithubSync.list_rag_files()
                        else:
                            st.warning(f"知识库重建失败: {rebuild_msg}")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("❌ 上传到GitHub失败")

    st.markdown("---")
    st.markdown("**🔧 手动维护：**")
    st.caption("如果自动重建未生效，可以手动触发从云端重建知识库。")
    if st.button("🔄 手动从云端重建知识库", key="manual_rebuild_kb"):
        with st.spinner("正在从云端拉取并重建知识库..."):
            success, msg = load_rag_from_github(aliyun_key)
            if success:
                st.success(msg)
                st.session_state.github_rag_files = GithubSync.list_rag_files()
            else:
                st.error(msg)
        time.sleep(1)
        st.rerun()


# ==========================================
# [TAB 4] 判例库设计（基础 + 进阶）
# ==========================================

with tab4:
    col_basic, col_supp = st.columns([5, 5])

    # ===== 左侧：基础判例 =====
    with col_basic:
        st.subheader("📗 基础判例")
        st.caption("最基础的判例将作为基本信息，全部提供给评分模型学习。")

        if st.button("📋 展示当前基础判例", use_container_width=True, key="show_basic"):
            show_basic_cases_dialog(embedder)

        # 编辑弹窗触发
        if st.session_state.get('editing_basic_idx') is not None:
            edit_basic_case_dialog(st.session_state.editing_basic_idx)

        # 手动添加
        with st.expander("➕ 手动添加基础判例"):
            with st.form("basic_case_form"):
                f_txt = st.text_area("判例描述", height=80, key="bc_txt")
                st.markdown("**因子评分详情**")
                fc1, fc2 = st.columns(2)
                factors = ["优雅性", "辨识度", "协调性", "饱和度", "持久性", "苦涩度"]
                input_scores = {}
                for i, f in enumerate(factors):
                    with (fc1 if i % 2 == 0 else fc2):
                        val = st.number_input(f"{f}分数", 0, 9, 7, key=f"bc_s_{i}")
                        cmt = st.text_input(f"{f}评语", key=f"bc_c_{i}")
                        sug = st.text_input(f"{f}建议", key=f"bc_a_{i}")
                        input_scores[f] = {"score": val, "comment": cmt, "suggestion": sug}
                f_master = st.text_area("宗师总评", key="bc_master", height=60)
                if st.form_submit_button("保存基础判例并同步"):
                    new_c = {
                        "text": f_txt, "scores": input_scores,
                        "master_comment": f_master, "created_at": time.strftime("%Y-%m-%d")
                    }
                    st.session_state.basic_cases.append(new_c)
                    ResourceManager.save_json(st.session_state.basic_cases, PATHS.basic_case_data)
                    with st.spinner("同步到GitHub..."):
                        GithubSync.sync_basic_cases(st.session_state.basic_cases)
                    st.success("已保存并同步！")
                    time.sleep(1)
                    st.rerun()

        # 批量添加
        with st.expander("📦 批量添加基础判例"):
            template_bytes = get_template_bytes()
            if template_bytes:
                st.download_button("📥 下载判例模板 (Excel)", template_bytes, "template.xlsx",
                                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                   key="dl_tpl_basic")
            else:
                st.warning("模板文件暂不可用")
            st.caption("请按照模板格式填写判例数据。如有多组判例，请在同一工作簿中新增 Sheet 页，每个 Sheet 页对应一条判例，格式与模板保持一致。")
            bc_file = st.file_uploader("上传已填写的判例文件", type=['xlsx', 'xls'], key="bc_batch_upload")
            if bc_file and st.button("📤 导入基础判例", key="bc_batch_import"):
                with st.spinner("正在处理文件..."):
                    new_cases = basic_case_process(bc_file)
                if new_cases:
                    st.session_state.basic_cases.extend(new_cases)
                    ResourceManager.save_json(st.session_state.basic_cases, PATHS.basic_case_data)
                    with st.spinner("同步到GitHub..."):
                        GithubSync.sync_basic_cases(st.session_state.basic_cases)
                    st.success(f"✅ 成功导入 {len(new_cases)} 条基础判例！")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("未能从文件中解析到有效判例，请检查文件格式。")

    # ===== 右侧：进阶判例 =====
    with col_supp:
        st.subheader("📘 进阶判例")
        st.caption("进阶判例将作为额外信息，经相似度比较后，筛选出与待评茶品最相似的部分提供给评分模型学习。")

        if st.button("📋 展示当前进阶判例", use_container_width=True, key="show_supp"):
            show_supp_cases_dialog(embedder)

        # 编辑弹窗触发
        if st.session_state.get('editing_supp_idx') is not None:
            edit_supp_case_dialog(st.session_state.editing_supp_idx, embedder)

        # 手动添加
        with st.expander("➕ 手动添加进阶判例"):
            with st.form("supp_case_form"):
                f_txt2 = st.text_area("判例描述", height=80, key="sc_txt")
                st.markdown("**因子评分详情**")
                fc1, fc2 = st.columns(2)
                input_scores2 = {}
                for i, f in enumerate(factors):
                    with (fc1 if i % 2 == 0 else fc2):
                        val = st.number_input(f"{f}分数", 0, 9, 7, key=f"sc_s_{i}")
                        cmt = st.text_input(f"{f}评语", key=f"sc_c_{i}")
                        sug = st.text_input(f"{f}建议", key=f"sc_a_{i}")
                        input_scores2[f] = {"score": val, "comment": cmt, "suggestion": sug}
                f_master2 = st.text_area("宗师总评", key="sc_master", height=60)
                if st.form_submit_button("保存进阶判例并同步"):
                    new_c = {
                        "text": f_txt2, "scores": input_scores2,
                        "master_comment": f_master2, "created_at": time.strftime("%Y-%m-%d")
                    }
                    supp_idx, supp_data = st.session_state.supp_cases
                    supp_data.append(new_c)
                    vec = embedder.encode([f_txt2])
                    supp_idx.add(vec)
                    st.session_state.supp_cases = (supp_idx, supp_data)
                    ResourceManager.save(supp_idx, supp_data, PATHS.supp_case_index, PATHS.supp_case_data, is_json=True)
                    with st.spinner("同步到GitHub..."):
                        GithubSync.sync_supp_cases(supp_data)
                    st.success("已保存并同步！")
                    time.sleep(1)
                    st.rerun()

        # 批量添加
        with st.expander("📦 批量添加进阶判例"):
            template_bytes2 = get_template_bytes()
            if template_bytes2:
                st.download_button("📥 下载判例模板 (Excel)", template_bytes2, "template.xlsx",
                                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                   key="dl_tpl_supp")
            else:
                st.warning("模板文件暂不可用")
            st.caption("请按照模板格式填写判例数据。如有多组判例，请在同一工作簿中新增 Sheet 页，每个 Sheet 页对应一条判例，格式与模板保持一致。")
            sc_file = st.file_uploader("上传已填写的判例文件", type=['xlsx', 'xls'], key="sc_batch_upload")
            if sc_file and st.button("📤 导入进阶判例", key="sc_batch_import"):
                with st.spinner("正在处理文件..."):
                    new_cases = supplementary_case_process(sc_file)
                if new_cases:
                    supp_idx, supp_data = st.session_state.supp_cases
                    supp_data.extend(new_cases)
                    # 重建 FAISS
                    new_idx = faiss.IndexFlatIP(1024)
                    vecs = embedder.encode([c["text"] for c in supp_data])
                    new_idx.add(vecs)
                    st.session_state.supp_cases = (new_idx, supp_data)
                    ResourceManager.save(new_idx, supp_data, PATHS.supp_case_index, PATHS.supp_case_data, is_json=True)
                    with st.spinner("同步到GitHub..."):
                        GithubSync.sync_supp_cases(supp_data)
                    st.success(f"✅ 成功导入 {len(new_cases)} 条进阶判例！")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("未能从文件中解析到有效判例，请检查文件格式。")


# ==========================================
# [TAB 5] 模型微调 (LoRA)
# ==========================================

with tab5:
    MANAGER_URL = "http://117.50.138.123:8001"
    st.subheader("🚀 模型微调 (LoRA)")

    # 检测服务器状态
    server_status = "unknown"
    try:
        resp = requests.get(f"{MANAGER_URL}/status", timeout=2)
        if resp.status_code == 200:
            status_data = resp.json()
            server_status = "idle" if status_data.get("vllm_status") == "running" else "training"
        else:
            server_status = "error"
    except:
        server_status = "offline"

    if server_status == "idle":
        st.success("🟢 服务器就绪（正在进行推理服务）")
    elif server_status == "training":
        st.warning("🟠 正在微调训练中...（推理服务暂停）")
        st.markdown("⚠️ **注意：** 此时无法进行评分交互，请耐心等待训练完成。")
    elif server_status == "offline":
        st.error("🔴 无法连接到 GPU 服务器（请联系管理员）")

    # 当前微调数据统计
    if PATHS.training_file.exists():
        with open(PATHS.training_file, "r", encoding="utf-8") as f:
            data_count = len(f.readlines())
    else:
        data_count = 0
    st.info(f"当前微调数据：**{data_count} 条** | 基础判例：**{len(st.session_state.basic_cases)}** 条 | 进阶判例：**{len(st.session_state.supp_cases[1])}** 条")

    ft_c1, ft_c2, ft_c3 = st.columns(3)

    # --- Column 1: 手动准备数据 ---
    with ft_c1:
        st.markdown("#### 📄 手动准备数据")
        template_bytes_ft = get_template_bytes()
        if template_bytes_ft:
            st.download_button("📥 下载微调数据模板 (Excel)", template_bytes_ft, "template.xlsx",
                               "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                               key="dl_tpl_ft")
        else:
            st.warning("模板文件暂不可用")
        st.caption("请按照模板格式填写微调训练数据。如有多组数据，请在同一工作簿中新增 Sheet 页，每个 Sheet 页对应一条训练样本，格式与模板保持一致。")
        ft_file = st.file_uploader("上传已填写的微调数据文件", type=['xlsx', 'xls'], key="ft_data_upload")
        if ft_file and st.button("📤 导入微调数据", key="ft_import"):
            with st.spinner("正在处理文件..."):
                new_entries = finetune_data_process(ft_file)
            if new_entries:
                # new_entries 应为可直接写入 JSONL 的 dict 列表
                try:
                    with open(PATHS.training_file, "a", encoding="utf-8") as f:
                        for entry in new_entries:
                            f.write(json.dumps(entry, ensure_ascii=False) + "\n")
                    st.success(f"✅ 成功导入 {len(new_entries)} 条微调数据！")
                    time.sleep(1)
                    st.rerun()
                except Exception as e:
                    st.error(f"写入失败: {e}")
            else:
                st.error("未能从文件中解析到有效数据，请检查文件格式。")

    # --- Column 2: 自动填充数据 ---
    with ft_c2:
        st.markdown("#### 🤖 自动填充数据")
        st.caption("从判例库自动提取数据并追加到微调文件（自动跳过重复数据）。")

        sys_tpl = st.session_state.prompt_config.get('system_template', '')
        user_tpl = st.session_state.prompt_config.get('user_template', '')

        if st.button("➕ 将基础判例添加至微调数据", use_container_width=True, key="ft_add_basic"):
            if not st.session_state.basic_cases:
                st.warning("基础判例库为空")
            else:
                added, skipped = ResourceManager.append_cases_to_finetune(
                    st.session_state.basic_cases, sys_tpl, user_tpl
                )
                st.success(f"新增 {added} 条，跳过 {skipped} 条重复数据")
                time.sleep(1)
                st.rerun()

        if st.button("➕ 将进阶判例添加至微调数据", use_container_width=True, key="ft_add_supp"):
            _, supp_data = st.session_state.supp_cases
            if not supp_data:
                st.warning("进阶判例库为空")
            else:
                added, skipped = ResourceManager.append_cases_to_finetune(
                    supp_data, sys_tpl, user_tpl
                )
                st.success(f"新增 {added} 条，跳过 {skipped} 条重复数据")
                time.sleep(1)
                st.rerun()

    # --- Column 3: 启动训练 ---
    with ft_c3:
        st.markdown("#### 🔥 启动训练")
        st.caption("点击下方按钮将把数据上传至 GPU 服务器并开始训练。训练期间推理服务将中断约 2-5 分钟。")
        btn_disabled = (server_status != "idle") or (data_count == 0)
        if st.button("🔥 开始微调 (Start LoRA)", type="primary", disabled=btn_disabled, key="start_ft"):
            if not PATHS.training_file.exists():
                st.error("找不到训练数据文件！")
            else:
                try:
                    with open(PATHS.training_file, "rb") as f:
                        with st.spinner("正在上传数据并启动训练任务..."):
                            files = {'file': ('tea_feedback.jsonl', f, 'application/json')}
                            r = requests.post(f"{MANAGER_URL}/upload_and_train", files=files, timeout=100)
                        if r.status_code == 200:
                            st.balloons()
                            st.success(f"✅ 任务已提交！服务器响应: {r.json().get('message')}")
                            st.info("💡 稍后刷新页面查看状态，训练完成后服务会自动恢复。")
                        else:
                            st.error(f"❌ 提交失败: {r.text}")
                except Exception as e:
                    st.error(f"❌ 连接错误: {e}")


# ==========================================
# [TAB 6] 提示词（Prompt）配置
# ==========================================

with tab6:
    pc = st.session_state.prompt_config
    st.markdown("系统提示词**可以修改**。完整全面的提示词会让大语言模型返回更准确的结果。")
    sys_t = st.text_area("系统提示词", pc.get('system_template', ''), height=350, key="sys_prompt_edit")

    # 保存按钮（位于用户提示词说明之前）
    save_col, restore_col = st.columns(2)
    with save_col:
        if st.button("💾 保存提示词修改", type="primary", key="save_prompt"):
            if sys_t == pc.get('system_template'):
                st.info("内容没有变化，无需保存。")
            else:
                new_cfg = {"system_template": sys_t, "user_template": pc.get('user_template', DEFAULT_USER_TEMPLATE)}
                with st.spinner("正在连接云端仓库并写入数据..."):
                    success = GithubSync.push_json(
                        file_path_in_repo="tea_data/prompts.json",
                        data_dict=new_cfg,
                        commit_msg="Update prompts.json from App"
                    )
                if success:
                    st.success("✅ 成功写入云端！")
                    st.session_state.prompt_config = new_cfg
                    with open(PATHS.prompt_config_file, 'w', encoding='utf-8') as f:
                        json.dump(new_cfg, f, ensure_ascii=False, indent=2)
                    time.sleep(1)
                    st.rerun()

    with restore_col:
        if st.button("🔄 将提示词恢复至初始设定", key="restore_prompt"):
            # 从 tea_backup/default_prompts.json 读取默认配置
            default_cfg = None
            if PATHS.default_prompts.exists():
                default_cfg = ResourceManager.load_external_json(PATHS.default_prompts)
            if not default_cfg:
                # 尝试从 GitHub 下载
                content = GithubSync.download_github_file("tea_backup/default_prompts.json")
                if content:
                    try:
                        default_cfg = json.loads(content.decode('utf-8'))
                    except:
                        pass
            if default_cfg and 'system_template' in default_cfg:
                st.session_state.prompt_config['system_template'] = default_cfg['system_template']
                # 同时保存到本地
                with open(PATHS.prompt_config_file, 'w', encoding='utf-8') as f:
                    json.dump(st.session_state.prompt_config, f, ensure_ascii=False, indent=2)
                st.success("✅ 已恢复至初始设定，请确认后点击「保存提示词修改」同步到云端。")
                time.sleep(1)
                st.rerun()
            else:
                st.error("❌ 未找到默认提示词配置文件 (tea_backup/default_prompts.json)")

    st.markdown("用户提示词**不可修改**。其保证了发送内容与回答内容的基本结构，因此大语言模型的回答可被准确解析。")
    st.text_area("用户提示词", pc.get('user_template', ''), height=250, disabled=True, key="user_prompt_view")





