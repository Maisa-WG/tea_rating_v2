"""
helpers.py
==========
辅助工具函数
"""

from typing import List, Dict, Optional
from io import BytesIO
from pathlib import Path

from PyPDF2 import PdfReader
from docx import Document


# ==========================================
# 文件解析函数
# ==========================================

def parse_file(uploaded_file) -> str:
    """
    解析上传文件为纯文本

    Args:
        uploaded_file: Streamlit UploadedFile 对象

    Returns:
        str: 解析后的文本内容
    """
    try:
        if uploaded_file.name.endswith('.txt'):
            return uploaded_file.read().decode("utf-8")
        if uploaded_file.name.endswith('.pdf'):
            return "".join([p.extract_text() for p in PdfReader(uploaded_file).pages])
        if uploaded_file.name.endswith('.docx'):
            return "\n".join([p.text for p in Document(uploaded_file).paragraphs])
    except:
        return ""
    return ""


def parse_file_bytes(filename: str, content: bytes) -> str:
    """
    解析文件内容（从 bytes）—— 用于从 GitHub 拉取的文件

    Args:
        filename: 文件名
        content: 文件字节内容

    Returns:
        str: 解析后的文本内容
    """
    try:
        if filename.lower().endswith('.txt'):
            return content.decode('utf-8', errors='ignore')
        elif filename.lower().endswith('.pdf'):
            if not content.startswith(b'%PDF'):
                print(f"[ERROR] 不是有效的 PDF 文件: {filename}")
                return ""
            reader = PdfReader(BytesIO(content))
            text = ""
            for page in reader.pages:
                try:
                    pt = page.extract_text()
                    if pt:
                        text += pt + "\n"
                except:
                    continue
            return text
        elif filename.lower().endswith('.docx'):
            doc = Document(BytesIO(content))
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        print(f"[ERROR] 解析 {filename} 失败: {e}")
    return ""


# ==========================================
# 报告生成函数
# ==========================================

def create_word_report(results: List[Dict]) -> BytesIO:
    """
    生成 Word 格式的批量评分报告

    Args:
        results: 评分结果列表

    Returns:
        BytesIO: Word 文档字节流
    """
    doc = Document()
    doc.add_heading("茶评批���评分报告", 0)
    for item in results:
        doc.add_heading(f"条目 {item['id']}", 1)
        doc.add_paragraph(f"原文：{item['text']}")
        s = item.get('scores', {}).get('scores', {})
        mc = item.get('scores', {}).get('master_comment', '')
        if mc:
            doc.add_paragraph(f"总评：{mc}", style="Intense Quote")
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '因子', '分数', '评语', '建议'
        for k, v in s.items():
            r = table.add_row().cells
            r[0].text = k
            r[1].text = str(v.get('score', ''))
            r[2].text = v.get('comment', '')
            r[3].text = v.get('suggestion', '')
        doc.add_paragraph("_" * 20)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ==========================================
# 文本分割函数（批量评分专用）
# ==========================================

def split_tea_reviews(text: str, max_length: int = 500) -> List[str]:
    """
    将长文本分割成多个独立的茶评条目

    Args:
        text: 原始文本
        max_length: 单条茶评的最大长度（字符数）

    Returns:
        List[str]: 分割后的茶评条目列表
    """
    # 清理文本：移除多余的空白字符
    text = text.strip()
    text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())

    # 如果文本很短，直接返回
    if len(text) <= max_length:
        return [text] if text else []

    reviews = []
    current_review = ""

    # 按行分割
    lines = text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 如果当前行以特定标记开头（如"茶品"、"样品"等），可能是新的茶评
        if line.startswith(('茶品', '样品', '品名', '【', '1.', '2.', '3.', '一、', '二、', '三、')):
            # 保存之前的茶评
            if current_review:
                reviews.append(current_review.strip())

            # 开始新的茶评
            current_review = line
        else:
            # 追加到当前茶评
            if current_review:
                current_review += " " + line
            else:
                current_review = line

        # 检查长度，如果超过最大长度，分割
        if len(current_review) >= max_length:
            reviews.append(current_review.strip())
            current_review = ""

    # 添加最后一个茶评
    if current_review:
        reviews.append(current_review.strip())

    # 如果没有成功分割，按段落分割
    if not reviews:
        paragraphs = text.split('\n\n')
        current = ""
        for para in paragraphs:
            if len(current) + len(para) <= max_length:
                current += para + "\n\n"
            else:
                if current:
                    reviews.append(current.strip())
                current = para + "\n\n"
        if current:
            reviews.append(current.strip())

    # 如果仍然没有结果，按固定长度强制分割
    if not reviews and len(text) > max_length:
        for i in range(0, len(text), max_length):
            reviews.append(text[i:i+max_length].strip())

    # 过滤空字符串
    reviews = [r for r in reviews if r and len(r) > 10]  # 至少10个字符

    return reviews


def parse_batch_file(uploaded_file) -> List[str]:
    """
    解析批量评分上传的文件，返回茶评条目列表

    Args:
        uploaded_file: Streamlit UploadedFile 对象

    Returns:
        List[str]: 茶评条目列表
    """
    try:
        # 解析文件内容
        text = parse_file(uploaded_file)

        if not text:
            return []

        # 分割成多个茶评
        reviews = split_tea_reviews(text)

        return reviews

    except Exception as e:
        print(f"[ERROR] 批量文件解析失败: {e}")
        return []


# ==========================================
# 模板下载函数
# ==========================================

def get_template_bytes(template_path: Path, github_sync_class) -> Optional[bytes]:
    """
    获取模板文件内容（优先本地，其次从 GitHub 下载）

    Args:
        template_path: 本地模板文件路径
        github_sync_class: GithubSync 类（用于从 GitHub 下载）

    Returns:
        Optional[bytes]: 模板文件内容，失败返回 None
    """
    if template_path.exists():
        with open(template_path, 'rb') as f:
            return f.read()
    # 尝试从 GitHub 下载
    content = github_sync_class.download_github_file("tea_backup/template.xlsx")
    if content:
        # 缓存到本地
        try:
            with open(template_path, 'wb') as f:
                f.write(content)
        except:
            pass
    return content